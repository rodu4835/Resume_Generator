import os
import re
import json
import docx
import openai
import warnings
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.text.run import Run
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from config import OPENAI_API_KEY

warnings.filterwarnings("ignore", category=UserWarning)

openai.api_key = OPENAI_API_KEY

# Function to read resume data from JSON file
def read_resume_from_json():
    with open('/home/ron/Programs/ResumeGenerator/data/resume.json', 'r') as f:
        return json.load(f)

def main(job_posting):
    resume = read_resume_from_json()
    desktop_path = "/path/to/Desktop/"
    if job_posting:
        new_resume = str(generateResume(job_posting, resume))
        # Extracting the JSON-like data from txt file
        merged_resume = extract_json(desktop_path, resume, new_resume)
        create_docx_from_json(merged_resume, desktop_path)
        cover_letter = generateCoverLetter(job_posting, merged_resume)
        create_cover_letter_docx(cover_letter, desktop_path)
    else:
        create_docx_from_json(resume, desktop_path)
    return "New resume saved!"
        
def generateResume(job_posting, resume):
    prompt = f"Take a look at my resume in JSON format seen here: {resume}"
    prompt += f"Based on the following job posting, {job_posting}\n"
    prompt += "Do not use my name, Only using the information from my existing Summary and Experience sections of my resume, generate a new summary and incorporate the Company name and Job title into my summary as a position I am interested in."
    prompt += "From the projects within my Experience section of the resume, identify which ones showcase my most relevant experience, pick 3. In order by most relevant, respond with the names of the 3 projects you've chosen within the Experience: tag."
    prompt += "Select the Technical and Soft skills that are relevant to the job and remove the rest, do not alter the programming languages."
    prompt += "Your response should be 4 seperate sections their respective titles, 'Summary':, 'Experience':, 'Technical': , 'Soft Skills':, with no other text but your changes."
    response = openai.ChatCompletion.create(
        model="gpt-4",  # Use the GPT-4 model or the desired chat model
        messages=[
            {"role": "system", "content": f"You are an AI assistant that will take in a job posting and resume and make the requested changes then respond in plain text format."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1000  # Set your desired max_tokens
    )
    return  response.choices[0].message["content"]

def generateCoverLetter(job_posting, resume):
    prompt = f"Take a look at my resume in JSON format seen here: {resume}"
    prompt += f"Based on the following job posting, {job_posting}\n"
    prompt += "Generate a cover letter for this job. It should be addressed to the Hiring Manager and end with my contact information."
    prompt += "Your response should have no other text except for the cover letter."
    response = openai.ChatCompletion.create(
        model="gpt-4",  # Use the GPT-4 model or the desired chat model
        messages=[
            {"role": "system", "content": f"You are an AI assistant that generates coverletters based on a job posting and resume."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1000  # Set your desired max_tokens
    )
    return  response.choices[0].message["content"]

#
def extract_json(desktop_path, resume, new_resume):
    new_resume = new_resume.replace("'", '"')
    new_resume = new_resume.replace("\\", "")
    new_resume = new_resume.replace('Bachelor"s',"Bachelor's")
    new_resume = "{" + new_resume + "}"
    
    file_path = os.path.join(desktop_path, "new_resume.txt")
    with open(file_path, 'w') as f:
        f.write(new_resume)
    
    new_summary = extractSummary(new_resume)
    new_technical = extractTechnical(new_resume)
    new_softSkills = extractSoftSkills(new_resume)
    new_experience = extractExperience(resume, new_resume)
    
    merged_resume = update_json(resume, new_summary, new_experience, new_technical, new_softSkills)
    
    # Write the merged resume back to a new JSON file
    merged_file_path = os.path.join(desktop_path, "merged_resume.json")
    with open(merged_file_path, 'w') as f:
        json.dump(merged_resume, f)
    return merged_resume

def update_json(resume_json, new_summary, selected_projects, new_technical, new_soft_skills):
    # Update Summary
    resume_json['Summary'] = new_summary
    
    # Replace Experience with Selected Projects
    resume_json['Experience'] = selected_projects
        
    # Navigate through the nested structure
    for skill_set in resume_json.get('Skills', []):
        # Update Technical
        if 'Technical' in skill_set:
            skill_set['Technical'] = new_technical
        # Update Soft Skills
        if 'Soft Skills' in skill_set:
            skill_set['Soft Skills'] = new_soft_skills
    
    return resume_json

def extractSummary(new_resume):
    # Use regex to find the summary string
    summary_match = re.search(r'"Summary":\s*"(.*?)"', new_resume, re.DOTALL)

    # Extract the summary string
    if summary_match:
        summary_string = summary_match.group(1)
    else:
        summary_string = "Summary not found"
    return summary_string

def extractExperience(resume, new_resume):
    common_projects = []
    
    project_names_json = [project['Name'] for project in resume['Experience']]
    
    # Search for project names in the text file using regex
    for project_name in project_names_json:
        if re.search(re.escape(project_name), new_resume, re.IGNORECASE):
            common_projects.append(project_name)
            
    relevant_projects = []
    
    # Loop through each project in the "Experience" section of the json_data
    for project in resume['Experience']:
        if project['Name'] in common_projects:
            relevant_projects.append(project)
    
    return relevant_projects


def extractTechnical(new_resume):
    # Use regex to find the technical skills string
    technical_match = re.search(r'"Technical":\s*(\[.*?\])', new_resume, re.DOTALL)
    
    # Extract the technical skills string
    if technical_match:
        technical_string = technical_match.group(1)
    else:
        technical_string = "Technical skills not found"
    
    # Optionally, you can convert this to a Python list
    try:
        technical_list = eval(technical_string)
    except:
        technical_list = "Failed to convert to list"
    return technical_list

def extractSoftSkills(new_resume):
    # Use regex to find the technical skills string
    softSkills_match = re.search(r'"Soft Skills":\s*(\[.*?\])', new_resume, re.DOTALL)
    
    # Extract the Soft Skills string
    if softSkills_match:
        softSkills_string = softSkills_match.group(1)
    else:
        softSkills_string = "Soft Skills not found"
    
    # Optionally, you can convert this to a Python list
    try:
        softSkills_list = eval(softSkills_string)
    except:
        softSkills_list = "Failed to convert to list"
    return softSkills_list

def create_docx_from_json(json_data, desktop_path):
    doc = Document()
    
    # Define font sizes
    heading_font_size = Pt(12)
    text_font_size = Pt(10)
    
    # Function to set font size for a paragraph
    def set_font_size(para, font_size):
        for run in para.runs:
            run.font.size = font_size
            
    # Function to set spacing for a paragraph
    def set_spacing(para):
        para.paragraph_format.space_before = Pt(5)
        para.paragraph_format.space_after = Pt(5)
        para.paragraph_format.line_spacing = Pt(12.15)  # Set line spacing to 1
    
    def set_font_size_for_run(run, font_size):
        run.font.size = font_size
        
    # Add 'Contact' as a new line, bold, and uppercase
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('CONTACT')
    run.bold = True
    run.underline = True
    set_font_size_for_run(run, heading_font_size)
    set_spacing(paragraph)
    
    # Add Contact without heading, formatted in one line
    contact = json_data['Contact'][0]
    # Add basic contact details
    contact_values = [
        contact['Name'],
        contact['Email'],
        contact['Phone'],
        contact['LinkedIn'],
        contact['GitHub']
    ]
    contact_line = " | ".join(contact_values)
    para = doc.add_paragraph(contact_line)
    run = para.runs[0]
    run.font.size = text_font_size
    set_spacing(para)
    

    # Add 'SUMMARY' as a new line, bold, and uppercase
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('SUMMARY')
    run.bold = True
    run.underline = True
    set_font_size_for_run(run, heading_font_size)
    set_spacing(paragraph)

    para = doc.add_paragraph(json_data['Summary'])
    set_font_size(para, text_font_size)
    set_spacing(para)

    
    # Add 'EXPERIENCE' as a new line, bold, and uppercase
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('EXPERIENCE')
    run.bold = True
    run.underline = True
    set_font_size_for_run(run, heading_font_size)
    set_spacing(paragraph)

    for exp in json_data['Experience']:
        para = doc.add_paragraph(exp['Name'])
        run = para.runs[0]
        run.font.size = text_font_size
        set_spacing(para)

        para = doc.add_paragraph(f"Technologies Used: {exp['Technologies Used']}", style='ListBullet')
        set_font_size(para, text_font_size)
        set_spacing(para)
        
        para = doc.add_paragraph(f"Description: {exp['Description']}", style='ListBullet')
        set_font_size(para, text_font_size)
        set_spacing(para)
        
        para = doc.add_paragraph(f"Skills Demonstrated: {exp['Skills Demonstrated']}", style='ListBullet')
        set_font_size(para, text_font_size)
        set_spacing(para)

    
    # Add 'EDUCATION' as a new line, bold, and uppercase
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('EDUCATION')
    run.bold = True
    run.underline = True
    set_font_size_for_run(run, heading_font_size)
    set_spacing(paragraph)

    for edu in json_data['Education']:
        para = doc.add_paragraph(f"{edu['Degree']}")
        run = para.runs[0]
        run.font.size = text_font_size
        set_spacing(para)

    
    # Add 'SKILLS' as a new line, bold, and uppercase
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('SKILLS')
    run.bold = True
    run.underline = True
    set_font_size_for_run(run, heading_font_size)
    set_spacing(paragraph)
    
    skills = json_data['Skills'][0]
    
    para = doc.add_paragraph("", style='ListBullet')
    para.add_run("Technical Skills: ")
    para.add_run(', '.join(skills['Technical']))
    set_font_size(para, text_font_size)
    set_spacing(para)
    
    para = doc.add_paragraph("", style='ListBullet')
    para.add_run("Programming Languages: ")
    para.add_run(', '.join(skills['Programming Languages']))
    set_font_size(para, text_font_size)
    set_spacing(para)
    
    para = doc.add_paragraph("", style='ListBullet')
    para.add_run("Soft Skills: ")
    para.add_run(', '.join(skills['Soft Skills']))
    set_font_size(para, text_font_size)
    set_spacing(para)

    # Save the document
    file_path_docx = os.path.join(desktop_path, "new_resume.docx")
    doc.save(file_path_docx)
    
def create_cover_letter_docx(cover_letter, desktop_path):
    doc = Document()
    # Add the string as a new paragraph
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(cover_letter)
    # Set the font size to 10pt
    font = run.font
    font.size = Pt(10)
    file_path = os.path.join(desktop_path, "new_cover_letter.docx")
    doc.save(file_path)
    
if __name__ == "__main__":
    app.run(debug=True)
